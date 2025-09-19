import os
os.environ["OPENAI_API_KEY"] = "" 

import logging
from mcp.server.fastmcp import FastMCP

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
from langchain.chains import RetrievalQA
from langchain_openai import ChatOpenAI, OpenAIEmbeddings

from docx import Document as WordDocument
import pandas as pd

logging.basicConfig(level=logging.INFO)

mcp = FastMCP("Office-RAG")

OFFICE_DIR = "C:\\vibeCoding\\rag-server\\office"

def load_office_documents(folder_path: str) -> list[Document]:
    docs = []

    for filename in os.listdir(folder_path):
        path = os.path.join(folder_path, filename)

        if filename.endswith(".docx"):
            word = WordDocument(path)
            full_text = "\n".join([p.text for p in word.paragraphs if p.text.strip()])
            docs.append(Document(page_content=full_text, metadata={"source": filename}))

        elif filename.endswith(".xlsx"):
            try:
                excel = pd.read_excel(path, sheet_name=None)
                for sheet_name, df in excel.items():
                    text = df.astype(str).to_string(index=False)
                    docs.append(Document(page_content=text, metadata={"source": f"{filename} - {sheet_name}"}))
            except Exception as e:
                logging.error(f"엑셀 파일 처리 오류: {filename} - {e}")

    return docs

raw_docs = load_office_documents(OFFICE_DIR)
splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
docs = splitter.split_documents(raw_docs)

embeddings = OpenAIEmbeddings()
llm = ChatOpenAI(model="gpt-4o")

vectorstore = Chroma.from_documents(
    documents=docs,
    embedding=embeddings
)

qa_chain = RetrievalQA.from_chain_type(
    llm=llm,
    retriever=vectorstore.as_retriever()
)

@mcp.tool()
def ask_office(query: str) -> str:
    """폴더 내 Word/Excel 문서를 기반으로 질문에 답변합니다."""
    logging.info(f"Query: {query}")
    try:
        return qa_chain.run(query)
    except Exception as e:
        logging.error(f"ask_office 실행 중 오류: {e}")
        return (
            "죄송합니다. ask_office 도구 실행 중 오류가 발생했습니다.\n"
            f"[에러 내용] {e}"
        )

if __name__ == "__main__":
    mcp.run(transport="stdio")
